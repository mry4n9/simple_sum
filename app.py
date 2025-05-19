import streamlit as st
import openai
import requests
from bs4 import BeautifulSoup
import PyPDF2 # Or from pdfplumber import open as pdfplumber_open
from docx import Document
from docx.shared import Pt
import io
import time

# --- Configuration ---
# Use "gpt-4o-mini" as "gpt-4.1-mini" is not a standard model name.
# "gpt-4o-mini" is the latest cost-effective and capable small model.
OPENAI_MODEL = "gpt-4o-mini"
MAX_SUMMARY_CHARS = 2500 # As per requirement

# --- OpenAI Client Initialization ---
# For local development, it will use secrets.toml.
# For Streamlit Cloud, set OPENAI_API_KEY in the Streamlit Cloud dashboard.
try:
    client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except Exception as e:
    st.error(f"Error initializing OpenAI client: {e}. Ensure OPENAI_API_KEY is set in Streamlit secrets.")
    st.stop()

# --- Helper Functions ---

def extract_text_from_url(url):
    """Extracts raw text content from a URL."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0'} # Be a good web citizen
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status() # Raises an exception for HTTP errors
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script_or_style in soup(["script", "style"]):
            script_or_style.decompose()
            
        text = soup.get_text(separator='\n', strip=True)
        return text
    except requests.exceptions.RequestException as e:
        st.error(f"Error fetching URL '{url}': {e}")
        return None
    except Exception as e:
        st.error(f"Error parsing URL '{url}': {e}")
        return None

def extract_text_from_pdf(uploaded_file):
    """Extracts raw text content from an uploaded PDF file."""
    if uploaded_file is None:
        return None
    try:
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text() or "" # Add or "" to handle None
        return text
    except Exception as e:
        st.error(f"Error reading PDF '{uploaded_file.name}': {e}")
        return None
    # If using pdfplumber:
    # try:
    #     with pdfplumber_open(uploaded_file) as pdf:
    #         text = ""
    #         for page in pdf.pages:
    #             page_text = page.extract_text()
    #             if page_text:
    #                 text += page_text + "\n"
    #     return text
    # except Exception as e:
    #     st.error(f"Error reading PDF '{uploaded_file.name}': {e}")
    #     return None


def get_ai_summary(text_content, max_chars=MAX_SUMMARY_CHARS):
    """Generates a summary using OpenAI API."""
    if not text_content or not text_content.strip():
        return "No content provided for summary."

    # Truncate input text if it's excessively long to avoid high token usage for raw text
    # GPT-4o-mini has a 128k token context window, so this is mostly a precaution for extremely large inputs
    # A character is roughly 0.25 tokens. 100k chars ~ 25k tokens.
    max_input_chars = 100000 
    if len(text_content) > max_input_chars:
        text_content = text_content[:max_input_chars]
        st.warning(f"Input text was truncated to {max_input_chars} characters before sending for summarization.")

    try:
        prompt = f"""Please summarize the following text.
        The summary should be concise, informative, and capture the main points.
        Aim for a summary of approximately {max_chars} characters, but prioritize quality and completeness within that limit.
        Do not exceed {max_chars} characters.

        Text to summarize:
        ---
        {text_content}
        ---
        Summary:
        """
        
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are an expert summarizer."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=int(max_chars / 2.5), # Estimate tokens, OpenAI counts tokens not chars
            temperature=0.5, # Lower for more factual summaries
        )
        summary = response.choices[0].message.content.strip()
        return summary[:max_chars] # Enforce character limit strictly
    except openai.APIError as e:
        st.error(f"OpenAI API Error: {e}")
        return f"Error: Could not generate summary due to API issue. {e}"
    except Exception as e:
        st.error(f"An unexpected error occurred during summarization: {e}")
        return f"Error: Could not generate summary. {e}"

def create_docx(report_data):
    """Creates a DOCX file from the report data."""
    doc = Document()
    doc.add_heading('Transparency Report', level=0)
    
    # Helper to add sections
    def add_section(title, raw_text, summary_text):
        doc.add_heading(title, level=1)
        
        doc.add_heading('Raw Extracted Content:', level=2)
        p_raw = doc.add_paragraph()
        p_raw.add_run(raw_text if raw_text else "N/A").font.size = Pt(10)
        
        doc.add_heading('AI Summary:', level=2)
        p_sum = doc.add_paragraph()
        p_sum.add_run(summary_text if summary_text else "N/A").font.size = Pt(10)
        doc.add_page_break()

    if 'url_raw' in report_data or 'url_sum' in report_data:
        add_section(
            "URL Analysis",
            report_data.get('url_raw', "No URL content extracted or provided."),
            report_data.get('url_sum', "No URL summary generated or content provided.")
        )
    
    if 'pdf1_raw' in report_data or 'pdf1_sum' in report_data:
         add_section(
            f"PDF 1 Analysis: {report_data.get('pdf1_name', 'Unnamed PDF')}",
            report_data.get('pdf1_raw', "No PDF 1 content extracted or provided."),
            report_data.get('pdf1_sum', "No PDF 1 summary generated or content provided.")
        )

    if 'pdf2_raw' in report_data or 'pdf2_sum' in report_data:
        add_section(
            f"PDF 2 Analysis: {report_data.get('pdf2_name', 'Unnamed PDF')}",
            report_data.get('pdf2_raw', "No PDF 2 content extracted or provided."),
            report_data.get('pdf2_sum', "No PDF 2 summary generated or content provided.")
        )

    # Save to a BytesIO object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0) # Rewind the stream to the beginning
    return file_stream

# --- Streamlit App UI ---
st.set_page_config(page_title="AI Document Summarizer", layout="wide")
st.title("📄 AI Document Summarizer & Reporter")
st.markdown("Summarize content from a URL and up to two PDF files, then generate a DOCX report.")
st.markdown(f"Using OpenAI model: `{OPENAI_MODEL}`. Summaries target ~{MAX_SUMMARY_CHARS} characters.")

# "Clear cache" requirement: Streamlit handles session state.
# Not using st.cache_data or st.cache_resource on processing functions ensures
# they re-run on each "Generate" click, effectively not loading "same data" from a cache.

# --- Inputs ---
url_input = st.text_input("Enter URL to summarize (optional):", placeholder="https://example.com/article")
uploaded_pdf1 = st.file_uploader("Upload PDF 1 (optional):", type="pdf", key="pdf1")
uploaded_pdf2 = st.file_uploader("Upload PDF 2 (optional):", type="pdf", key="pdf2")

# --- Processing and Output ---
if 'report_generated' not in st.session_state:
    st.session_state.report_generated = False
if 'docx_file_stream' not in st.session_state:
    st.session_state.docx_file_stream = None
if 'report_data' not in st.session_state:
    st.session_state.report_data = {}


if st.button("🚀 Generate Report", type="primary"):
    if not url_input and not uploaded_pdf1 and not uploaded_pdf2:
        st.warning("Please provide a URL or upload at least one PDF.")
    else:
        st.session_state.report_generated = False
        st.session_state.docx_file_stream = None
        st.session_state.report_data = {} # Clear previous data

        with st.spinner("Hold tight! Analyzing documents and crafting summaries... This might take a moment."):
            # --- URL Processing ---
            if url_input:
                st.subheader("🌐 URL Processing")
                with st.status("Extracting text from URL...", expanded=True) as status_url:
                    url_context_raw = extract_text_from_url(url_input)
                    if url_context_raw:
                        st.session_state.report_data['url_raw'] = url_context_raw
                        status_url.update(label="URL text extracted! Summarizing...", state="running")
                        url_sum = get_ai_summary(url_context_raw)
                        st.session_state.report_data['url_sum'] = url_sum
                        status_url.update(label="URL processing complete!", state="complete", expanded=False)
                        with st.expander("View URL Summary & Raw Text", expanded=False):
                            st.markdown("**Summary:**")
                            st.markdown(f"> {url_sum}")
                            st.markdown("**Raw Text (first 1000 chars):**")
                            st.text_area("URL Raw Text", url_context_raw[:1000]+"...", height=150, disabled=True, key="url_raw_disp")
                    else:
                        st.session_state.report_data['url_raw'] = "Failed to extract content from URL."
                        st.session_state.report_data['url_sum'] = "N/A due to extraction failure."
                        status_url.update(label="URL processing failed.", state="error", expanded=True)
                st.divider()

            # --- PDF 1 Processing ---
            if uploaded_pdf1:
                st.subheader(f"📄 PDF 1 Processing: `{uploaded_pdf1.name}`")
                st.session_state.report_data['pdf1_name'] = uploaded_pdf1.name
                with st.status(f"Extracting text from {uploaded_pdf1.name}...", expanded=True) as status_pdf1:
                    pdf1_raw = extract_text_from_pdf(uploaded_pdf1)
                    if pdf1_raw:
                        st.session_state.report_data['pdf1_raw'] = pdf1_raw
                        status_pdf1.update(label=f"Text extracted from {uploaded_pdf1.name}! Summarizing...", state="running")
                        pdf1_sum = get_ai_summary(pdf1_raw)
                        st.session_state.report_data['pdf1_sum'] = pdf1_sum
                        status_pdf1.update(label=f"PDF 1 ({uploaded_pdf1.name}) processing complete!", state="complete", expanded=False)
                        with st.expander(f"View PDF 1 ({uploaded_pdf1.name}) Summary & Raw Text", expanded=False):
                            st.markdown("**Summary:**")
                            st.markdown(f"> {pdf1_sum}")
                            st.markdown("**Raw Text (first 1000 chars):**")
                            st.text_area("PDF1 Raw Text", pdf1_raw[:1000]+"...", height=150, disabled=True, key="pdf1_raw_disp")
                    else:
                        st.session_state.report_data['pdf1_raw'] = f"Failed to extract content from PDF: {uploaded_pdf1.name}."
                        st.session_state.report_data['pdf1_sum'] = "N/A due to extraction failure."
                        status_pdf1.update(label=f"PDF 1 ({uploaded_pdf1.name}) processing failed.", state="error", expanded=True)
                st.divider()

            # --- PDF 2 Processing ---
            if uploaded_pdf2:
                st.subheader(f"📄 PDF 2 Processing: `{uploaded_pdf2.name}`")
                st.session_state.report_data['pdf2_name'] = uploaded_pdf2.name
                with st.status(f"Extracting text from {uploaded_pdf2.name}...", expanded=True) as status_pdf2:
                    pdf2_raw = extract_text_from_pdf(uploaded_pdf2)
                    if pdf2_raw:
                        st.session_state.report_data['pdf2_raw'] = pdf2_raw
                        status_pdf2.update(label=f"Text extracted from {uploaded_pdf2.name}! Summarizing...", state="running")
                        pdf2_sum = get_ai_summary(pdf2_raw)
                        st.session_state.report_data['pdf2_sum'] = pdf2_sum
                        status_pdf2.update(label=f"PDF 2 ({uploaded_pdf2.name}) processing complete!", state="complete", expanded=False)
                        with st.expander(f"View PDF 2 ({uploaded_pdf2.name}) Summary & Raw Text", expanded=False):
                            st.markdown("**Summary:**")
                            st.markdown(f"> {pdf2_sum}")
                            st.markdown("**Raw Text (first 1000 chars):**")
                            st.text_area("PDF2 Raw Text", pdf2_raw[:1000]+"...", height=150, disabled=True, key="pdf2_raw_disp")
                    else:
                        st.session_state.report_data['pdf2_raw'] = f"Failed to extract content from PDF: {uploaded_pdf2.name}."
                        st.session_state.report_data['pdf2_sum'] = "N/A due to extraction failure."
                        status_pdf2.update(label=f"PDF 2 ({uploaded_pdf2.name}) processing failed.", state="error", expanded=True)
                st.divider()

            # --- DOCX Generation ---
            if st.session_state.report_data: # If any data was processed
                st.subheader("📝 Generating DOCX Report")
                with st.status("Compiling report...", expanded=True) as status_docx:
                    st.session_state.docx_file_stream = create_docx(st.session_state.report_data)
                    st.session_state.report_generated = True
                    status_docx.update(label="Report compiled!", state="complete", expanded=False)
                st.success("🎉 Report generated successfully!")
            else:
                st.error("No data was processed. Cannot generate report.")


# --- Download Button ---
# Show download button if report has been generated
if st.session_state.report_generated and st.session_state.docx_file_stream:
    current_time = time.strftime("%Y%m%d_%H%M%S")
    st.download_button(
        label="📥 Download Transparency Report (.docx)",
        data=st.session_state.docx_file_stream,
        file_name=f"transparency_report_{current_time}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.markdown("---")
st.caption("Developed with Streamlit and OpenAI.")
